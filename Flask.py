import os
import re
import fitz
import xlwt
from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
from docx import Document as DocxDocument
from spire.doc import *

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    text = ""
    with fitz.open(pdf_path) as pdf_file:
        for page_num in range(len(pdf_file)):
            page = pdf_file.load_page(page_num)
            text += page.get_text()
    return text


def extract_text_from_docx(docx_path):
    """Extracts text from a Docx file."""
    doc = DocxDocument(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    # Remove the evaluation warning text if present
    text = re.sub(r'Evaluation Warning: The document was created with Spire\.Doc for Python\.', '', text)
    return text


def convert_doc_to_docx(file_path):
    try:
        # Convert .doc file to .docx using Spire.Doc
        document = Document()
        document.LoadFromFile(file_path)
        docx_path = file_path.replace('.doc', '.docx')
        document.SaveToFile(docx_path, FileFormat.Docx2016)
        document.Close()

        print(f"Converted {file_path} to .docx")
        return docx_path
    except Exception as e:
        print(f"Error converting .doc to .docx: {e}")
        return None


def parse_resume_from_file(file_path):
    """Parses a resume file and extracts key information."""
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        text = extract_text_from_docx(file_path)
    elif file_extension == '.doc':
        # Convert .doc to .docx if necessary
        docx_path = convert_doc_to_docx(file_path)
        if docx_path:
            text = extract_text_from_docx(docx_path)
        else:
            return None, None, None
    else:
        return None, None, None

    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    email = email_match.group(0) if email_match else None

    # Improved phone number extraction
    phone_match = re.search(r'(\+\d{1,3}[\s-]?)?\(?\d{3}\)?[\s.-]?\d{2,5}[\s.-]?\d{4}', text)
    phone = phone_match.group(0) if phone_match else None

    # Extract all text from the file.
    all_text = text.strip()

    # Return the extracted information.
    return email, phone, all_text


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'files[]' not in request.files:
            return render_template('index.html', error='No file part')
        files = request.files.getlist('files[]')
        if len(files) == 0:
            return render_template('index.html', error='No selected files')
        parsed_resumes = []
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)

                try:
                    email, phone, all_text = parse_resume_from_file(file_path)
                    if email:
                        parsed_resumes.append((email, phone, all_text))
                except Exception as e:
                    return render_template('index.html', error=f'Error processing file {filename}: {e}')
        return write_resumes_to_excel(parsed_resumes)
    return render_template('index.html')


import os

def write_resumes_to_excel(parsed_resumes):
    """Writes parsed resumes to an Excel file."""
    output_folder = os.path.dirname(os.path.abspath(__file__))  # Get the directory of the Flask script
    output_file = os.path.join(output_folder, 'parsed_resumes.xls')  # Output file path

    workbook = xlwt.Workbook()
    worksheet = workbook.add_sheet('Resumes')

    # Write headers
    headers = ['Email', 'Phone', 'All Text']
    for col, header in enumerate(headers):
        worksheet.write(0, col, header)

    # Write data
    for row, resume in enumerate(parsed_resumes, start=1):
        for col, value in enumerate(resume):
            worksheet.write(row, col, value)

    # Save the workbook
    workbook.save(output_file)

    return send_file(output_file, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
